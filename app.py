import os
from datetime import date, datetime
from flask import (Flask, render_template, request,
                   redirect, url_for, flash, send_file)
from flask_login import (LoginManager, login_user, logout_user,
                         login_required, current_user)
from werkzeug.security import generate_password_hash
from database import conectar, inicializar_banco, DATABASE
import sqlite3
import shutil
from functools import wraps

from auth import carregar_usuario, verificar_login

# Importações para PDF e DOC
from io import BytesIO
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "escola_biblica_2026")

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"
login_manager.login_message = "Faça login para continuar."
login_manager.login_message_category = "warning"

@login_manager.user_loader
def load_user(user_id):
    return carregar_usuario(user_id)

inicializar_banco()

# Decorador para exigir perfil de administrador
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or current_user.perfil != 'admin':
            flash("Acesso negado. Apenas administradores podem acessar esta página.", "erro")
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function


# ══════════════════════════════════════
# AUTH
# ══════════════════════════════════════
@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        return redirect(url_for("index"))
    if request.method == "POST":
        email   = request.form.get("email", "").strip()
        senha   = request.form.get("senha", "")
        usuario = verificar_login(email, senha)
        if usuario:
            login_user(usuario, remember=True)
            flash(f"Bem-vindo(a), {usuario.nome}!", "sucesso")
            return redirect(
                request.args.get("next") or url_for("index"))
        flash("E-mail ou senha incorretos!", "erro")
    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    nome = current_user.nome
    logout_user()
    flash(f"Até logo, {nome}!", "sucesso")
    return redirect(url_for("login"))


# ══════════════════════════════════════
# PAINEL
# ══════════════════════════════════════
@app.route("/")
@login_required
def index():
    conn   = conectar()
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) as t FROM alunos")
    total_alunos = cursor.fetchone()["t"]
    cursor.execute("SELECT COUNT(*) as t FROM professores")
    total_professores = cursor.fetchone()["t"]
    cursor.execute(
        "SELECT COUNT(*) as t FROM disciplinas WHERE ativa=1")
    total_disciplinas = cursor.fetchone()["t"]
    cursor.execute(
        "SELECT COUNT(*) as t FROM turmas WHERE ativa=1")
    total_turmas = cursor.fetchone()["t"]
    cursor.execute(
        "SELECT COUNT(*) as t FROM matriculas WHERE status='aprovado'")
    aprovados = cursor.fetchone()["t"]
    cursor.execute(
        "SELECT COUNT(*) as t FROM matriculas WHERE status='reprovado'")
    reprovados = cursor.fetchone()["t"]
    cursor.execute(
        "SELECT COUNT(*) as t FROM matriculas WHERE status='cursando'")
    cursando = cursor.fetchone()["t"]
    conn.close()
    return render_template("index.html",
        total_alunos=total_alunos,
        total_professores=total_professores,
        total_disciplinas=total_disciplinas,
        total_turmas=total_turmas,
        aprovados=aprovados,
        reprovados=reprovados,
        cursando=cursando)


# ══════════════════════════════════════
# TURMAS
# ══════════════════════════════════════
@app.route("/turmas")
@login_required
def turmas():
    conn   = conectar()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT t.*, COUNT(a.id) as total_alunos
        FROM turmas t
        LEFT JOIN alunos a ON a.turma_id = t.id
        GROUP BY t.id ORDER BY t.nome
    """)
    lista = cursor.fetchall()
    conn.close()
    return render_template("turmas.html", turmas=lista)


@app.route("/turmas/nova", methods=["GET", "POST"])
@login_required
def nova_turma():
    if request.method == "POST":
        nome         = request.form.get("nome", "").strip()
        descricao    = request.form.get("descricao", "").strip()
        faixa_etaria = request.form.get("faixa_etaria", "").strip()
        if not nome:
            flash("Nome é obrigatório!", "erro")
            return redirect(url_for("nova_turma"))
        conn   = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute(
                "INSERT INTO turmas (nome,descricao,faixa_etaria) VALUES (?,?,?)",
                (nome, descricao, faixa_etaria))
            conn.commit()
            flash(f"Turma '{nome}' criada!", "sucesso")
        except sqlite3.IntegrityError as e:
            if "turmas.nome" in str(e):
                flash("Já existe uma turma com este nome!", "erro")
            else:
                flash(f"Erro de integridade ao cadastrar turma: {e}", "erro")
        except Exception as e:
            flash(f"Erro inesperado ao cadastrar turma: {e}", "erro")
        finally:
            conn.close()
        return redirect(url_for("turmas"))
    return render_template("nova_turma.html")


@app.route("/turmas/<int:id>/editar", methods=["GET", "POST"])
@login_required
def editar_turma(id):
    conn   = conectar()
    cursor = conn.cursor()
    if request.method == "POST":
        nome         = request.form.get("nome", "").strip()
        descricao    = request.form.get("descricao", "").strip()
        faixa_etaria = request.form.get("faixa_etaria", "").strip()
        ativa        = 1 if request.form.get("ativa") == "on" else 0
        if not nome:
            flash("Nome é obrigatório!", "erro")
            conn.close()
            return redirect(url_for("editar_turma", id=id))
        try:
            cursor.execute(
                "UPDATE turmas SET nome=?, descricao=?, faixa_etaria=?, ativa=? WHERE id=?",
                (nome, descricao, faixa_etaria, ativa, id))
            conn.commit()
            flash(f"Turma '{nome}' atualizada!", "sucesso")
        except sqlite3.IntegrityError as e:
            if "turmas.nome" in str(e):
                flash("Já existe uma turma com este nome!", "erro")
            else:
                flash(f"Erro de integridade ao atualizar turma: {e}", "erro")
        except Exception as e:
            flash(f"Erro inesperado ao atualizar turma: {e}", "erro")
        finally:
            conn.close()
        return redirect(url_for("turmas"))
    else:
        cursor.execute("SELECT * FROM turmas WHERE id=?", (id,))
        turma = cursor.fetchone()
        conn.close()
        if turma:
            return render_template("editar_turma.html", turma=turma)
        flash("Turma não encontrada!", "erro")
        return redirect(url_for("turmas"))


@app.route("/turmas/<int:id>/excluir", methods=["POST"])
@login_required
def excluir_turma(id):
    conn   = conectar()
    cursor = conn.cursor()
    try:
        # Verificar se há alunos associados a esta turma
        cursor.execute("SELECT COUNT(*) FROM alunos WHERE turma_id = ?", (id,))
        total_alunos = cursor.fetchone()[0]
        if total_alunos > 0:
            flash(f"Não é possível excluir a turma. Há {total_alunos} aluno(s) associado(s) a ela.", "erro")
            return redirect(url_for("turmas"))

        cursor.execute("DELETE FROM turmas WHERE id=?", (id,))
        conn.commit()
        flash("Turma excluída!", "sucesso")
    except Exception as e:
        flash(f"Erro ao excluir turma: {e}", "erro")
    finally:
        conn.close()
    return redirect(url_for("turmas"))


# ══════════════════════════════════════
# ALUNOS
# ══════════════════════════════════════
@app.route("/alunos")
@login_required
def alunos():
    conn   = conectar()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT a.*, t.nome as turma_nome, t.faixa_etaria
        FROM alunos a
        LEFT JOIN turmas t ON a.turma_id = t.id
        ORDER BY a.nome
    """)
    lista = cursor.fetchall()
    conn.close()
    return render_template("alunos.html", alunos=lista)


@app.route("/alunos/novo", methods=["GET", "POST"])
@login_required
def novo_aluno():
    conn = conectar()
    cursor = conn.cursor()
    turmas = []
    try:
        cursor.execute("SELECT id, nome, faixa_etaria FROM turmas WHERE ativa=1 ORDER BY nome")
        turmas = cursor.fetchall()
    except Exception as e:
        flash(f"Erro ao carregar turmas: {e}", "erro")
    finally:
        conn.close()

    if request.method == "POST":
        nome            = request.form.get("nome", "").strip()
        data_nascimento = request.form.get("data_nascimento", "").strip() or None
        telefone        = request.form.get("telefone", "").strip() or None
        email           = request.form.get("email", "").strip() or None
        membro_igreja   = 1 if request.form.get("membro_igreja") == "on" else 0
        turma_id        = request.form.get("turma_id", type=int) or None
        # Novos campos
        nome_pai        = request.form.get("nome_pai", "").strip() or None
        nome_mae        = request.form.get("nome_mae", "").strip() or None
        endereco        = request.form.get("endereco", "").strip() or None

        if not nome:
            flash("Nome do aluno é obrigatório!", "erro")
            # Recarregar turmas para o template em caso de erro
            conn = conectar()
            cursor = conn.cursor()
            cursor.execute("SELECT id, nome, faixa_etaria FROM turmas WHERE ativa=1 ORDER BY nome")
            turmas = cursor.fetchall()
            conn.close()
            return render_template("novo_aluno.html", turmas=turmas)

        conn = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                INSERT INTO alunos (nome, data_nascimento, telefone, email, membro_igreja, turma_id, nome_pai, nome_mae, endereco)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (nome, data_nascimento, telefone, email, membro_igreja, turma_id, nome_pai, nome_mae, endereco))
            conn.commit()
            flash(f"Aluno '{nome}' cadastrado!", "sucesso")
        except Exception as e:
            flash(f"Erro ao cadastrar aluno: {e}", "erro")
            # Recarregar turmas para o template em caso de erro
            conn = conectar()
            cursor = conn.cursor()
            cursor.execute("SELECT id, nome, faixa_etaria FROM turmas WHERE ativa=1 ORDER BY nome")
            turmas = cursor.fetchall()
            conn.close()
            return render_template("novo_aluno.html", turmas=turmas)
        finally:
            conn.close()
        return redirect(url_for("alunos"))
    return render_template("novo_aluno.html", turmas=turmas)


@app.route("/alunos/<int:id>/editar", methods=["GET", "POST"])
@login_required
def editar_aluno(id):
    conn = conectar()
    cursor = conn.cursor()
    aluno = None
    turmas = []
    try:
        cursor.execute("SELECT * FROM alunos WHERE id=?", (id,))
        aluno = cursor.fetchone()
        cursor.execute("SELECT id, nome, faixa_etaria FROM turmas WHERE ativa=1 ORDER BY nome")
        turmas = cursor.fetchall()
    except Exception as e:
        flash(f"Erro ao carregar dados: {e}", "erro")
    finally:
        conn.close()

    if not aluno:
        flash("Aluno não encontrado!", "erro")
        return redirect(url_for("alunos"))

    if request.method == "POST":
        nome            = request.form.get("nome", "").strip()
        data_nascimento = request.form.get("data_nascimento", "").strip() or None
        telefone        = request.form.get("telefone", "").strip() or None
        email           = request.form.get("email", "").strip() or None
        membro_igreja   = 1 if request.form.get("membro_igreja") == "on" else 0
        turma_id        = request.form.get("turma_id", type=int) or None
        # Novos campos
        nome_pai        = request.form.get("nome_pai", "").strip() or None
        nome_mae        = request.form.get("nome_mae", "").strip() or None
        endereco        = request.form.get("endereco", "").strip() or None

        if not nome:
            flash("Nome do aluno é obrigatório!", "erro")
            return render_template("editar_aluno.html", aluno=aluno, turmas=turmas)

        conn = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                UPDATE alunos SET
                    nome=?, data_nascimento=?, telefone=?, email=?, membro_igreja=?, turma_id=?,
                    nome_pai=?, nome_mae=?, endereco=?
                WHERE id=?
            """, (nome, data_nascimento, telefone, email, membro_igreja, turma_id,
                  nome_pai, nome_mae, endereco, id))
            conn.commit()
            flash(f"Aluno '{nome}' atualizado!", "sucesso")
        except Exception as e:
            flash(f"Erro ao atualizar aluno: {e}", "erro")
            return render_template("editar_aluno.html", aluno=aluno, turmas=turmas)
        finally:
            conn.close()
        return redirect(url_for("alunos"))
    return render_template("editar_aluno.html", aluno=aluno, turmas=turmas)


@app.route("/alunos/<int:id>/excluir", methods=["POST"])
@login_required
def excluir_aluno(id):
    conn   = conectar()
    cursor = conn.cursor()
    try:
        # Verificar se o aluno tem matrículas ativas
        cursor.execute("SELECT COUNT(*) FROM matriculas WHERE aluno_id = ?", (id,))
        total_matriculas = cursor.fetchone()[0]
        if total_matriculas > 0:
            flash(f"Não é possível excluir o aluno. Há {total_matriculas} matrícula(s) associada(s) a ele.", "erro")
            return redirect(url_for("alunos"))

        cursor.execute("DELETE FROM alunos WHERE id=?", (id,))
        conn.commit()
        flash("Aluno excluído!", "sucesso")
    except Exception as e:
        flash(f"Erro ao excluir aluno: {e}", "erro")
    finally:
        conn.close()
    return redirect(url_for("alunos"))


@app.route("/alunos/<int:id>/trilha")
@login_required
def trilha_aluno(id):
    conn = conectar()
    cursor = conn.cursor()
    aluno_dict = None
    matriculas_processadas = []

    try:
        # Buscar dados do aluno
        cursor.execute("""
            SELECT a.*, t.nome as turma_nome, t.faixa_etaria
            FROM alunos a
            LEFT JOIN turmas t ON a.turma_id = t.id
            WHERE a.id = ?
        """, (id,))
        aluno = cursor.fetchone()
        if aluno:
            aluno_dict = dict(aluno) # Converter para dict mutável

            # Buscar matrículas do aluno
            cursor.execute("""
                SELECT m.id AS matricula_id,
                       d.nome AS disciplina_nome,
                       d.tem_atividades,
                       d.frequencia_minima,
                       m.data_inicio, m.data_conclusao, m.status,
                       m.nota1, m.nota2, m.participacao, m.desafio, m.prova,
                       m.meditacao, m.versiculos, m.desafio_nota, m.visitante
                FROM matriculas m
                JOIN disciplinas d ON m.disciplina_id = d.id
                WHERE m.aluno_id = ?
                ORDER BY m.data_inicio DESC
            """, (id,))
            raw_matriculas = cursor.fetchall()

            for mat in raw_matriculas:
                matricula_dict = dict(mat) # Converter para dict mutável

                # --- Cálculo de Frequência ---
                cursor.execute("""
                    SELECT data_aula, presente, fez_atividade
                    FROM presencas
                    WHERE matricula_id = ?
                    ORDER BY data_aula DESC
                """, (matricula_dict['matricula_id'],))
                historico_chamadas = [dict(c) for c in cursor.fetchall()] # Converter para dict

                presencas = sum(1 for c in historico_chamadas if c['presente'])
                total_aulas = len(historico_chamadas)
                atividades_feitas = sum(1 for c in historico_chamadas if c['fez_atividade'])

                matricula_dict['presencas'] = presencas
                matricula_dict['total_aulas'] = total_aulas
                matricula_dict['atividades_feitas'] = atividades_feitas
                matricula_dict['historico_chamadas'] = historico_chamadas # Adicionar histórico

                frequencia_porcentagem = (presencas / total_aulas * 100) if total_aulas > 0 else 0
                matricula_dict['frequencia_porcentagem'] = frequencia_porcentagem
                matricula_dict['frequencia_minima_atingida'] = frequencia_porcentagem >= matricula_dict['frequencia_minima']

                # --- Cálculo de Notas e Status ---
                nota_final_calc = None
                status_frequencia = None # Para crianças

                faixa_etaria = aluno_dict.get('faixa_etaria', 'adultos') # Usar a faixa etária do aluno

                if 'criancas' in faixa_etaria:
                    matricula_dict['media_display'] = 'N/A'
                    matricula_dict['tem_atividades'] = False # Crianças não têm atividades para nota
                    # Status de crianças baseado apenas na frequência
                    if matricula_dict['status'] == 'cursando':
                        status_frequencia = 'Aprovado (Provisório)' if matricula_dict['frequencia_minima_atingida'] else 'Reprovado (Provisório)'
                    elif matricula_dict['status'] == 'aprovado':
                        status_frequencia = 'Aprovado'
                    elif matricula_dict['status'] == 'reprovado':
                        status_frequencia = 'Reprovado'
                    else:
                        status_frequencia = matricula_dict['status'].capitalize()
                    matricula_dict['status_display'] = status_frequencia

                elif 'adolescentes' in faixa_etaria or 'jovens' in faixa_etaria:
                    meditacao = matricula_dict.get('meditacao') or 0
                    versiculos = matricula_dict.get('versiculos') or 0
                    desafio_nota = matricula_dict.get('desafio_nota') or 0
                    visitante = matricula_dict.get('visitante') or 0
                    nota_final_calc = meditacao + versiculos + desafio_nota + visitante
                    matricula_dict['media_display'] = f"{nota_final_calc:.1f}"

                    # Lógica de status para Adolescentes/Jovens
                    if matricula_dict['status'] == 'cursando':
                        if nota_final_calc >= 7.0 and matricula_dict['frequencia_minima_atingida']:
                            matricula_dict['status_display'] = 'Aprovado (Provisório)'
                        elif nota_final_calc < 7.0 and matricula_dict['frequencia_minima_atingida']:
                            matricula_dict['status_display'] = 'Reprovado por Nota (Provisório)'
                        elif nota_final_calc >= 7.0 and not matricula_dict['frequencia_minima_atingida']:
                            matricula_dict['status_display'] = 'Reprovado por Frequência (Provisório)'
                        else:
                            matricula_dict['status_display'] = 'Reprovado (Provisório)'
                    else:
                        matricula_dict['status_display'] = matricula_dict['status'].capitalize()

                else: # Adultos
                    nota1 = matricula_dict.get('nota1') or 0
                    nota2 = matricula_dict.get('nota2') or 0
                    participacao = matricula_dict.get('participacao') or 0
                    desafio = matricula_dict.get('desafio') or 0
                    prova = matricula_dict.get('prova') or 0

                    # Calcula N1 (participacao + desafio + prova)
                    n1_calculada = (participacao or 0) + (desafio or 0) + (prova or 0)
                    # Se nota1 foi preenchida, usa ela, senão usa a calculada
                    nota1_final = nota1 if matricula_dict.get('nota1') is not None else n1_calculada

                    # Média final é a média de nota1_final e nota2
                    if matricula_dict.get('nota2') is not None:
                        nota_final_calc = (nota1_final + nota2) / 2
                        matricula_dict['media_display'] = f"{nota_final_calc:.1f}"
                    else:
                        matricula_dict['media_display'] = f"{nota1_final:.1f}" # Se N2 não foi lançada, mostra N1

                    # Lógica de status para Adultos
                    if matricula_dict['status'] == 'cursando':
                        if nota_final_calc is not None and nota_final_calc >= 7.0 and matricula_dict['frequencia_minima_atingida']:
                            matricula_dict['status_display'] = 'Aprovado (Provisório)'
                        elif nota_final_calc is not None and nota_final_calc < 7.0 and matricula_dict['frequencia_minima_atingida']:
                            matricula_dict['status_display'] = 'Reprovado por Nota (Provisório)'
                        elif nota_final_calc is not None and nota_final_calc >= 7.0 and not matricula_dict['frequencia_minima_atingida']:
                            matricula_dict['status_display'] = 'Reprovado por Frequência (Provisório)'
                        else:
                            matricula_dict['status_display'] = 'Reprovado (Provisório)'
                    else:
                        matricula_dict['status_display'] = matricula_dict['status'].capitalize()

                matriculas_processadas.append(matricula_dict)

    except Exception as e:
        flash(f"Erro ao carregar trilha do aluno: {e}", "erro")
        print(f"ERRO NA TRILHA DO ALUNO: {e}") # Logar o erro no console
        return redirect(url_for('alunos')) # Redirecionar para a lista de alunos em caso de erro grave
    finally:
        conn.close()

    if aluno_dict:
        return render_template("trilha_aluno.html", aluno=aluno_dict, matriculas=matriculas_processadas)
    else:
        flash("Aluno não encontrado!", "erro")
        return redirect(url_for('alunos'))


# ══════════════════════════════════════
# DISCIPLINAS
# ══════════════════════════════════════
@app.route("/disciplinas")
@login_required
def disciplinas():
    conn   = conectar()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT d.*, u.nome as professor_nome, COUNT(m.id) as total_matriculas
        FROM disciplinas d
        LEFT JOIN usuarios u ON d.professor_id = u.id
        LEFT JOIN matriculas m ON m.disciplina_id = d.id
        GROUP BY d.id
        ORDER BY d.nome
    """)
    lista = cursor.fetchall()
    conn.close()
    return render_template("disciplinas.html", disciplinas=lista)


@app.route("/disciplinas/nova", methods=["GET", "POST"])
@login_required
def nova_disciplina():
    conn   = conectar()
    cursor = conn.cursor()
    professores = []
    try:
        cursor.execute("SELECT id, nome FROM usuarios WHERE perfil='professor' OR perfil='admin' ORDER BY nome")
        professores = cursor.fetchall()
    except Exception as e:
        flash(f"Erro ao carregar professores: {e}", "erro")
    finally:
        conn.close()

    if request.method == "POST":
        nome            = request.form.get("nome", "").strip()
        descricao       = request.form.get("descricao", "").strip()
        professor_id    = request.form.get("professor_id", type=int) or None
        tem_atividades  = 1 if request.form.get("tem_atividades") == "on" else 0
        frequencia_minima = request.form.get("frequencia_minima", type=float) or 75.0

        if not nome:
            flash("Nome da disciplina é obrigatório!", "erro")
            return render_template("nova_disciplina.html", professores=professores)

        conn = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                INSERT INTO disciplinas (nome, descricao, professor_id, tem_atividades, frequencia_minima)
                VALUES (?, ?, ?, ?, ?)
            """, (nome, descricao, professor_id, tem_atividades, frequencia_minima))
            conn.commit()
            flash(f"Disciplina '{nome}' criada!", "sucesso")
        except sqlite3.IntegrityError as e:
            if "disciplinas.nome" in str(e):
                flash("Já existe uma disciplina com este nome!", "erro")
            else:
                flash(f"Erro de integridade ao cadastrar disciplina: {e}", "erro")
        except Exception as e:
            flash(f"Erro ao cadastrar disciplina: {e}", "erro")
        finally:
            conn.close()
        return redirect(url_for("disciplinas"))
    return render_template("nova_disciplina.html", professores=professores)


@app.route("/disciplinas/<int:id>/editar", methods=["GET", "POST"])
@login_required
def editar_disciplina(id):
    conn   = conectar()
    cursor = conn.cursor()
    disciplina = None
    professores = []
    try:
        cursor.execute("SELECT * FROM disciplinas WHERE id=?", (id,))
        disciplina = cursor.fetchone()
        cursor.execute("SELECT id, nome FROM usuarios WHERE perfil='professor' OR perfil='admin' ORDER BY nome")
        professores = cursor.fetchall()
    except Exception as e:
        flash(f"Erro ao carregar dados: {e}", "erro")
    finally:
        conn.close()

    if not disciplina:
        flash("Disciplina não encontrada!", "erro")
        return redirect(url_for("disciplinas"))

    if request.method == "POST":
        nome            = request.form.get("nome", "").strip()
        descricao       = request.form.get("descricao", "").strip()
        professor_id    = request.form.get("professor_id", type=int) or None
        tem_atividades  = 1 if request.form.get("tem_atividades") == "on" else 0
        frequencia_minima = request.form.get("frequencia_minima", type=float) or 75.0
        ativa           = 1 if request.form.get("ativa") == "on" else 0

        if not nome:
            flash("Nome da disciplina é obrigatório!", "erro")
            return render_template("editar_disciplina.html", disciplina=disciplina, professores=professores)

        conn = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                UPDATE disciplinas SET
                    nome=?, descricao=?, professor_id=?, tem_atividades=?, frequencia_minima=?, ativa=?
                WHERE id=?
            """, (nome, descricao, professor_id, tem_atividades, frequencia_minima, ativa, id))
            conn.commit()
            flash(f"Disciplina '{nome}' atualizada!", "sucesso")
        except sqlite3.IntegrityError as e:
            if "disciplinas.nome" in str(e):
                flash("Já existe uma disciplina com este nome!", "erro")
            else:
                flash(f"Erro de integridade ao atualizar disciplina: {e}", "erro")
        except Exception as e:
            flash(f"Erro ao atualizar disciplina: {e}", "erro")
        finally:
            conn.close()
        return redirect(url_for("disciplinas"))
    return render_template("editar_disciplina.html", disciplina=disciplina, professores=professores)


@app.route("/disciplinas/<int:id>/excluir", methods=["POST"])
@login_required
def excluir_disciplina(id):
    conn   = conectar()
    cursor = conn.cursor()
    try:
        # Verificar se há matrículas associadas a esta disciplina
        cursor.execute("SELECT COUNT(*) FROM matriculas WHERE disciplina_id = ?", (id,))
        total_matriculas = cursor.fetchone()[0]
        if total_matriculas > 0:
            flash(f"Não é possível excluir a disciplina. Há {total_matriculas} matrícula(s) associada(s) a ela.", "erro")
            return redirect(url_for("disciplinas"))

        cursor.execute("DELETE FROM disciplinas WHERE id=?", (id,))
        conn.commit()
        flash("Disciplina excluída!", "sucesso")
    except Exception as e:
        flash(f"Erro ao excluir disciplina: {e}", "erro")
    finally:
        conn.close()
    return redirect(url_for("disciplinas"))


# ══════════════════════════════════════
# MATRÍCULAS
# ══════════════════════════════════════
@app.route("/matriculas")
@login_required
def matriculas():
    conn   = conectar()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT m.id, a.nome as aluno_nome, d.nome as disciplina_nome,
               m.data_inicio, m.data_conclusao, m.status,
               t.faixa_etaria, d.tem_atividades, d.frequencia_minima,
               m.nota1, m.nota2, m.participacao, m.desafio, m.prova,
               m.meditacao, m.versiculos, m.desafio_nota, m.visitante
        FROM matriculas m
        JOIN alunos a ON m.aluno_id = a.id
        JOIN disciplinas d ON m.disciplina_id = d.id
        LEFT JOIN turmas t ON a.turma_id = t.id
        ORDER BY a.nome, d.nome
    """)
    raw_matriculas = cursor.fetchall()
    processed_matriculas = []

    for mat in raw_matriculas:
        matricula_dict = dict(mat) # Converter para dict mutável

        # --- Cálculo de Frequência ---
        cursor.execute("""
            SELECT presente
            FROM presencas
            WHERE matricula_id = ?
        """, (matricula_dict['id'],))
        historico_chamadas = cursor.fetchall()

        presencas = sum(1 for c in historico_chamadas if c['presente'])
        total_aulas = len(historico_chamadas)

        matricula_dict['presencas'] = presencas
        matricula_dict['total_aulas'] = total_aulas

        frequencia_porcentagem = (presencas / total_aulas * 100) if total_aulas > 0 else 0
        matricula_dict['frequencia_porcentagem'] = frequencia_porcentagem
        matricula_dict['frequencia_minima_atingida'] = frequencia_porcentagem >= matricula_dict['frequencia_minima']

        # --- Cálculo de Notas e Status ---
        nota_final_calc = None
        faixa_etaria = matricula_dict.get('faixa_etaria', 'adultos')

        if 'criancas' in faixa_etaria:
            matricula_dict['media_display'] = 'N/A'
            matricula_dict['status_display'] = 'Aprovado (Frequência)' if matricula_dict['frequencia_minima_atingida'] else 'Reprovado (Frequência)'
        elif 'adolescentes' in faixa_etaria or 'jovens' in faixa_etaria:
            meditacao = matricula_dict.get('meditacao') or 0
            versiculos = matricula_dict.get('versiculos') or 0
            desafio_nota = matricula_dict.get('desafio_nota') or 0
            visitante = matricula_dict.get('visitante') or 0
            nota_final_calc = meditacao + versiculos + desafio_nota + visitante
            matricula_dict['media_display'] = f"{nota_final_calc:.1f}"
            if nota_final_calc >= 7.0 and matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Aprovado'
            elif nota_final_calc < 7.0 and matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Reprovado por Nota'
            elif nota_final_calc >= 7.0 and not matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Reprovado por Frequência'
            else:
                matricula_dict['status_display'] = 'Reprovado'
        else: # Adultos
            nota1 = matricula_dict.get('nota1') or 0
            nota2 = matricula_dict.get('nota2') or 0
            participacao = matricula_dict.get('participacao') or 0
            desafio = matricula_dict.get('desafio') or 0
            prova = matricula_dict.get('prova') or 0

            n1_calculada = (participacao or 0) + (desafio or 0) + (prova or 0)
            nota1_final = nota1 if matricula_dict.get('nota1') is not None else n1_calculada

            if matricula_dict.get('nota2') is not None:
                nota_final_calc = (nota1_final + nota2) / 2
                matricula_dict['media_display'] = f"{nota_final_calc:.1f}"
            else:
                matricula_dict['media_display'] = f"{nota1_final:.1f}"

            if nota_final_calc is not None and nota_final_calc >= 7.0 and matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Aprovado'
            elif nota_final_calc is not None and nota_final_calc < 7.0 and matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Reprovado por Nota'
            elif nota_final_calc is not None and nota_final_calc >= 7.0 and not matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Reprovado por Frequência'
            else:
                matricula_dict['status_display'] = 'Reprovado'

        processed_matriculas.append(matricula_dict)

    conn.close()
    return render_template("matriculas.html", matriculas=processed_matriculas)


@app.route("/matriculas/nova", methods=["GET", "POST"])
@login_required
def nova_matricula():
    conn = conectar()
    cursor = conn.cursor()
    alunos = []
    disciplinas = []
    try:
        cursor.execute("SELECT id, nome FROM alunos ORDER BY nome")
        alunos = cursor.fetchall()
        cursor.execute("SELECT id, nome FROM disciplinas WHERE ativa=1 ORDER BY nome")
        disciplinas = cursor.fetchall()
    except Exception as e:
        flash(f"Erro ao carregar dados: {e}", "erro")
    finally:
        conn.close()

    if request.method == "POST":
        aluno_id      = request.form.get("aluno_id", type=int)
        disciplina_id = request.form.get("disciplina_id", type=int)
        data_inicio   = request.form.get("data_inicio", "").strip()

        if not aluno_id or not disciplina_id or not data_inicio:
            flash("Todos os campos são obrigatórios!", "erro")
            return render_template("nova_matricula.html", alunos=alunos, disciplinas=disciplinas)

        conn = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute("""
                INSERT INTO matriculas (aluno_id, disciplina_id, data_inicio)
                VALUES (?, ?, ?)
            """, (aluno_id, disciplina_id, data_inicio))
            conn.commit()
            flash("Matrícula criada com sucesso!", "sucesso")
        except sqlite3.IntegrityError:
            flash("Este aluno já está matriculado nesta disciplina!", "erro")
        except Exception as e:
            flash(f"Erro ao criar matrícula: {e}", "erro")
        finally:
            conn.close()
        return redirect(url_for("matriculas"))
    return render_template("nova_matricula.html", alunos=alunos, disciplinas=disciplinas)


@app.route("/matriculas/novo_aluno_disciplina", methods=["GET", "POST"])
@login_required
def novo_aluno_disciplina():
    conn = conectar()
    cursor = conn.cursor()
    disciplinas = []
    turmas = []
    try:
        cursor.execute("SELECT id, nome FROM disciplinas WHERE ativa=1 ORDER BY nome")
        disciplinas = cursor.fetchall()
        cursor.execute("SELECT id, nome, faixa_etaria FROM turmas WHERE ativa=1 ORDER BY nome")
        turmas = cursor.fetchall()
    except Exception as e:
        flash(f"Erro ao carregar dados: {e}", "erro")
    finally:
        conn.close()

    if request.method == "POST":
        # Dados do Aluno
        nome_aluno      = request.form.get("nome_aluno", "").strip()
        data_nascimento = request.form.get("data_nascimento", "").strip() or None
        telefone        = request.form.get("telefone", "").strip() or None
        email           = request.form.get("email", "").strip() or None
        membro_igreja   = 1 if request.form.get("membro_igreja") == "on" else 0
        turma_id        = request.form.get("turma_id", type=int) or None
        nome_pai        = request.form.get("nome_pai", "").strip() or None
        nome_mae        = request.form.get("nome_mae", "").strip() or None
        endereco        = request.form.get("endereco", "").strip() or None

        # Dados da Matrícula
        disciplina_id   = request.form.get("disciplina_id", type=int)
        data_inicio     = request.form.get("data_inicio", "").strip()

        if not nome_aluno or not disciplina_id or not data_inicio:
            flash("Nome do aluno, disciplina e data de início são obrigatórios!", "erro")
            return render_template("novo_aluno_disciplina.html", disciplinas=disciplinas, turmas=turmas)

        conn = conectar()
        cursor = conn.cursor()
        try:
            # 1. Cadastrar o novo aluno
            cursor.execute("""
                INSERT INTO alunos (nome, data_nascimento, telefone, email, membro_igreja, turma_id, nome_pai, nome_mae, endereco)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (nome_aluno, data_nascimento, telefone, email, membro_igreja, turma_id, nome_pai, nome_mae, endereco))
            aluno_id = cursor.lastrowid

            # 2. Matricular o aluno na disciplina
            cursor.execute("""
                INSERT INTO matriculas (aluno_id, disciplina_id, data_inicio)
                VALUES (?, ?, ?)
            """, (aluno_id, disciplina_id, data_inicio))
            conn.commit()
            flash(f"Aluno '{nome_aluno}' cadastrado e matriculado em '{disciplina_id}' com sucesso!", "sucesso")
        except sqlite3.IntegrityError:
            flash("Erro: Possível duplicidade de aluno ou matrícula.", "erro")
        except Exception as e:
            flash(f"Erro ao cadastrar aluno e matrícula: {e}", "erro")
        finally:
            conn.close()
        return redirect(url_for("matriculas"))
    return render_template("novo_aluno_disciplina.html", disciplinas=disciplinas, turmas=turmas)


@app.route("/matriculas/<int:id>/editar", methods=["GET", "POST"])
@login_required
def editar_matricula(id):
    conn = conectar()
    cursor = conn.cursor()
    matricula = None
    aluno = None
    disciplina = None
    faixa_etaria = 'adultos' # Default

    try:
        cursor.execute("""
            SELECT m.*, a.nome as aluno_nome, d.nome as disciplina_nome, t.faixa_etaria
            FROM matriculas m
            JOIN alunos a ON m.aluno_id = a.id
            JOIN disciplinas d ON m.disciplina_id = d.id
            LEFT JOIN turmas t ON a.turma_id = t.id
            WHERE m.id = ?
        """, (id,))
        matricula = cursor.fetchone()
        if matricula:
            aluno = {'nome': matricula['aluno_nome']} # Apenas o nome é necessário para exibição
            disciplina = {'nome': matricula['disciplina_nome']} # Apenas o nome é necessário para exibição
            faixa_etaria = matricula['faixa_etaria'] if matricula['faixa_etaria'] else 'adultos'
    except Exception as e:
        flash(f"Erro ao carregar dados da matrícula: {e}", "erro")
    finally:
        conn.close()

    if not matricula:
        flash("Matrícula não encontrada!", "erro")
        return redirect(url_for("matriculas"))

    if request.method == "POST":
        data_inicio    = request.form.get("data_inicio", "").strip()
        data_conclusao = request.form.get("data_conclusao", "").strip() or None
        status         = request.form.get("status", "").strip()

        # Campos de notas para Adultos
        nota1_adulto = request.form.get("nota1_adulto", type=float)
        nota2_adulto = request.form.get("nota2_adulto", type=float)
        participacao_adulto = request.form.get("participacao_adulto", type=float)
        desafio_adulto = request.form.get("desafio_adulto", type=float)
        prova_adulto = request.form.get("prova_adulto", type=float)

        # Campos de notas para Adolescentes/Jovens
        meditacao_aj = request.form.get("meditacao_aj", type=float)
        versiculos_aj = request.form.get("versiculos_aj", type=float)
        desafio_nota_aj = request.form.get("desafio_nota_aj", type=float)
        visitante_aj = request.form.get("visitante_aj", type=float)

        conn = conectar()
        cursor = conn.cursor()
        try:
            # Resetar todas as notas para NULL antes de aplicar as específicas
            update_query = """
                UPDATE matriculas SET
                    data_inicio=?, data_conclusao=?, status=?,
                    nota1=NULL, nota2=NULL, participacao=NULL, desafio=NULL, prova=NULL,
                    meditacao=NULL, versiculos=NULL, desafio_nota=NULL, visitante=NULL
                WHERE id=?
            """
            cursor.execute(update_query, (data_inicio, data_conclusao, status, id))

            # Aplicar notas com base na faixa etária
            if 'criancas' in faixa_etaria:
                # Nenhuma nota para crianças, já resetado para NULL
                pass
            elif 'adolescentes' in faixa_etaria or 'jovens' in faixa_etaria:
                cursor.execute("""
                    UPDATE matriculas SET
                        meditacao=?, versiculos=?, desafio_nota=?, visitante=?
                    WHERE id=?
                """, (meditacao_aj, versiculos_aj, desafio_nota_aj, visitante_aj, id))
            else: # Adultos
                cursor.execute("""
                    UPDATE matriculas SET
                        nota1=?, nota2=?, participacao=?, desafio=?, prova=?
                    WHERE id=?
                """, (nota1_adulto, nota2_adulto, participacao_adulto, desafio_adulto, prova_adulto, id))

            conn.commit()
            flash("Matrícula atualizada com sucesso!", "sucesso")
        except Exception as e:
            flash(f"Erro ao atualizar matrícula: {e}", "erro")
        finally:
            conn.close()
        return redirect(url_for("matriculas"))
    return render_template("editar_matricula.html", matricula=matricula, aluno=aluno, disciplina=disciplina, faixa_etaria=faixa_etaria)


@app.route("/matriculas/<int:id>/excluir", methods=["POST"])
@login_required
def excluir_matricula(id):
    conn   = conectar()
    cursor = conn.cursor()
    try:
        # Excluir presenças associadas a esta matrícula primeiro
        cursor.execute("DELETE FROM presencas WHERE matricula_id = ?", (id,))
        cursor.execute("DELETE FROM matriculas WHERE id=?", (id,))
        conn.commit()
        flash("Matrícula excluída!", "sucesso")
    except Exception as e:
        flash(f"Erro ao excluir matrícula: {e}", "erro")
    finally:
        conn.close()
    return redirect(url_for("matriculas"))


# ══════════════════════════════════════
# PRESENÇA
# ══════════════════════════════════════
@app.route("/presenca/chamada", methods=["GET", "POST"])
@login_required
def chamada():
    conn = conectar()
    cursor = conn.cursor()
    disciplinas = []
    alunos_chamada = []
    selected_disciplina = None
    selected_data_aula = None
    tem_atividades = False

    try:
        cursor.execute("SELECT id, nome FROM disciplinas WHERE ativa=1 ORDER BY nome")
        disciplinas = cursor.fetchall()

        if request.method == "POST":
            disciplina_id = request.form.get("disciplina_id", type=int)
            data_aula_str = request.form.get("data_aula", "").strip()

            if not disciplina_id or not data_aula_str:
                flash("Selecione uma disciplina e uma data para a aula.", "erro")
                return render_template("chamada.html", disciplinas=disciplinas)

            selected_data_aula = data_aula_str
            cursor.execute("SELECT * FROM disciplinas WHERE id = ?", (disciplina_id,))
            selected_disciplina = cursor.fetchone()
            if selected_disciplina:
                tem_atividades = selected_disciplina['tem_atividades'] == 1

            # Buscar alunos matriculados na disciplina
            cursor.execute("""
                SELECT m.id AS matricula_id, a.nome AS aluno_nome,
                       p.presente, p.fez_atividade
                FROM matriculas m
                JOIN alunos a ON m.aluno_id = a.id
                LEFT JOIN presencas p ON p.matricula_id = m.id AND p.data_aula = ?
                WHERE m.disciplina_id = ? AND m.status = 'cursando'
                ORDER BY a.nome
            """, (data_aula_str, disciplina_id))
            alunos_chamada = [dict(row) for row in cursor.fetchall()] # Converter para dict

            # Atualizar presenças no banco de dados
            for aluno_data in alunos_chamada:
                matricula_id = aluno_data['matricula_id']
                presente = 1 if request.form.get(f"presente_{matricula_id}") == "on" else 0
                fez_atividade = 1 if request.form.get(f"atividade_{matricula_id}") == "on" else 0

                cursor.execute("""
                    INSERT INTO presencas (matricula_id, data_aula, presente, fez_atividade)
                    VALUES (?, ?, ?, ?)
                    ON CONFLICT(matricula_id, data_aula) DO UPDATE SET
                        presente = EXCLUDED.presente,
                        fez_atividade = EXCLUDED.fez_atividade
                """, (matricula_id, data_aula_str, presente, fez_atividade))
            conn.commit()
            flash("Chamada salva com sucesso!", "sucesso")

            # Recarregar alunos_chamada após salvar para refletir o estado atual
            cursor.execute("""
                SELECT m.id AS matricula_id, a.nome AS aluno_nome,
                       p.presente, p.fez_atividade
                FROM matriculas m
                JOIN alunos a ON m.aluno_id = a.id
                LEFT JOIN presencas p ON p.matricula_id = m.id AND p.data_aula = ?
                WHERE m.disciplina_id = ? AND m.status = 'cursando'
                ORDER BY a.nome
            """, (data_aula_str, disciplina_id))
            alunos_chamada = [dict(row) for row in cursor.fetchall()]

        elif request.method == "GET":
            disciplina_id = request.args.get("disciplina_id", type=int)
            data_aula_str = request.args.get("data_aula", "").strip()

            if disciplina_id and data_aula_str:
                selected_data_aula = data_aula_str
                cursor.execute("SELECT * FROM disciplinas WHERE id = ?", (disciplina_id,))
                selected_disciplina = cursor.fetchone()
                if selected_disciplina:
                    tem_atividades = selected_disciplina['tem_atividades'] == 1

                cursor.execute("""
                    SELECT m.id AS matricula_id, a.nome AS aluno_nome,
                           p.presente, p.fez_atividade
                    FROM matriculas m
                    JOIN alunos a ON m.aluno_id = a.id
                    LEFT JOIN presencas p ON p.matricula_id = m.id AND p.data_aula = ?
                    WHERE m.disciplina_id = ? AND m.status = 'cursando'
                    ORDER BY a.nome
                """, (data_aula_str, disciplina_id))
                alunos_chamada = [dict(row) for row in cursor.fetchall()]

    except Exception as e:
        flash(f"Erro ao carregar ou salvar chamada: {e}", "erro")
        print(f"ERRO NA CHAMADA: {e}") # Logar o erro no console
    finally:
        conn.close()

    return render_template("chamada.html",
                           disciplinas=disciplinas,
                           alunos_chamada=alunos_chamada,
                           selected_disciplina=selected_disciplina,
                           selected_data_aula=selected_data_aula,
                           tem_atividades=tem_atividades,
                           today=date.today())


# ══════════════════════════════════════
# RELATÓRIOS
# ══════════════════════════════════════
def get_relatorio_data(disciplina_id=None, turma_id=None, aluno_id=None, status=None):
    conn = conectar()
    cursor = conn.cursor()

    query = """
        SELECT m.id AS matricula_id,
               a.nome AS aluno_nome,
               d.nome AS disciplina_nome,
               t.nome AS turma_nome,
               t.faixa_etaria AS turma_faixa_etaria,
               m.data_inicio, m.data_conclusao, m.status,
               m.nota1, m.nota2, m.participacao, m.desafio, m.prova,
               m.meditacao, m.versiculos, m.desafio_nota, m.visitante,
               d.tem_atividades, d.frequencia_minima
        FROM matriculas m
        JOIN alunos a ON m.aluno_id = a.id
        JOIN disciplinas d ON m.disciplina_id = d.id
        LEFT JOIN turmas t ON a.turma_id = t.id
        WHERE 1=1
    """
    params = []

    if disciplina_id:
        query += " AND m.disciplina_id = ?"
        params.append(disciplina_id)
    if turma_id:
        query += " AND t.id = ?"
        params.append(turma_id)
    if aluno_id:
        query += " AND a.id = ?"
        params.append(aluno_id)
    if status:
        query += " AND m.status = ?"
        params.append(status)

    query += " ORDER BY a.nome, d.nome"
    cursor.execute(query, params)
    raw_matriculas = cursor.fetchall()
    processed_matriculas = []

    for mat in raw_matriculas:
        matricula_dict = dict(mat) # Converter para dict mutável

        # --- Cálculo de Frequência ---
        cursor.execute("""
            SELECT presente, fez_atividade
            FROM presencas
            WHERE matricula_id = ?
        """, (matricula_dict['matricula_id'],))
        historico_chamadas = cursor.fetchall()

        presencas = sum(1 for c in historico_chamadas if c['presente'])
        total_aulas = len(historico_chamadas)
        atividades_feitas = sum(1 for c in historico_chamadas if c['fez_atividade'])

        matricula_dict['presencas'] = presencas
        matricula_dict['total_aulas'] = total_aulas
        matricula_dict['atividades_feitas'] = atividades_feitas

        frequencia_porcentagem = (presencas / total_aulas * 100) if total_aulas > 0 else 0
        matricula_dict['frequencia_porcentagem'] = frequencia_porcentagem
        matricula_dict['frequencia_minima_atingida'] = frequencia_porcentagem >= matricula_dict['frequencia_minima']

        # --- Cálculo de Notas e Status ---
        nota_final_calc = None
        faixa_etaria = matricula_dict.get('turma_faixa_etaria', 'adultos')

        if 'criancas' in faixa_etaria:
            matricula_dict['media_display'] = 'N/A'
            matricula_dict['status_display'] = 'Aprovado (Frequência)' if matricula_dict['frequencia_minima_atingida'] else 'Reprovado (Frequência)'
        elif 'adolescentes' in faixa_etaria or 'jovens' in faixa_etaria:
            meditacao = matricula_dict.get('meditacao') or 0
            versiculos = matricula_dict.get('versiculos') or 0
            desafio_nota = matricula_dict.get('desafio_nota') or 0
            visitante = matricula_dict.get('visitante') or 0
            nota_final_calc = meditacao + versiculos + desafio_nota + visitante
            matricula_dict['media_display'] = f"{nota_final_calc:.1f}"
            if nota_final_calc >= 7.0 and matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Aprovado'
            elif nota_final_calc < 7.0 and matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Reprovado por Nota'
            elif nota_final_calc >= 7.0 and not matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Reprovado por Frequência'
            else:
                matricula_dict['status_display'] = 'Reprovado'
        else: # Adultos
            nota1 = matricula_dict.get('nota1') or 0
            nota2 = matricula_dict.get('nota2') or 0
            participacao = matricula_dict.get('participacao') or 0
            desafio = matricula_dict.get('desafio') or 0
            prova = matricula_dict.get('prova') or 0

            n1_calculada = (participacao or 0) + (desafio or 0) + (prova or 0)
            nota1_final = nota1 if matricula_dict.get('nota1') is not None else n1_calculada

            if matricula_dict.get('nota2') is not None:
                nota_final_calc = (nota1_final + nota2) / 2
                matricula_dict['media_display'] = f"{nota_final_calc:.1f}"
            else:
                matricula_dict['media_display'] = f"{nota1_final:.1f}"

            if nota_final_calc is not None and nota_final_calc >= 7.0 and matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Aprovado'
            elif nota_final_calc is not None and nota_final_calc < 7.0 and matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Reprovado por Nota'
            elif nota_final_calc is not None and nota_final_calc >= 7.0 and not matricula_dict['frequencia_minima_atingida']:
                matricula_dict['status_display'] = 'Reprovado por Frequência'
            else:
                matricula_dict['status_display'] = 'Reprovado'

        processed_matriculas.append(matricula_dict)

    conn.close()
    return processed_matriculas


@app.route("/relatorios")
@login_required
def relatorios():
    conn = conectar()
    cursor = conn.cursor()
    disciplinas = []
    turmas = []
    alunos = []
    try:
        cursor.execute("SELECT id, nome FROM disciplinas WHERE ativa=1 ORDER BY nome")
        disciplinas = cursor.fetchall()
        cursor.execute("SELECT id, nome FROM turmas WHERE ativa=1 ORDER BY nome")
        turmas = cursor.fetchall()
        cursor.execute("SELECT id, nome FROM alunos ORDER BY nome")
        alunos = cursor.fetchall()
    except Exception as e:
        flash(f"Erro ao carregar filtros: {e}", "erro")
    finally:
        conn.close()

    return render_template("relatorios.html",
                           disciplinas=disciplinas,
                           turmas=turmas,
                           alunos=alunos,
                           matriculas=[]) # Inicialmente vazio


@app.route("/relatorios/gerar", methods=["POST"])
@login_required
def gerar_relatorio():
    disciplina_id = request.form.get("disciplina_id", type=int) or None
    turma_id = request.form.get("turma_id", type=int) or None
    aluno_id = request.form.get("aluno_id", type=int) or None
    status = request.form.get("status", "").strip() or None

    matriculas = get_relatorio_data(disciplina_id, turma_id, aluno_id, status)

    conn = conectar()
    cursor = conn.cursor()
    disciplinas = []
    turmas = []
    alunos = []
    try:
        cursor.execute("SELECT id, nome FROM disciplinas WHERE ativa=1 ORDER BY nome")
        disciplinas = cursor.fetchall()
        cursor.execute("SELECT id, nome FROM turmas WHERE ativa=1 ORDER BY nome")
        turmas = cursor.fetchall()
        cursor.execute("SELECT id, nome FROM alunos ORDER BY nome")
        alunos = cursor.fetchall()
    except Exception as e:
        flash(f"Erro ao recarregar filtros: {e}", "erro")
    finally:
        conn.close()

    return render_template("relatorios.html",
                           disciplinas=disciplinas,
                           turmas=turmas,
                           alunos=alunos,
                           matriculas=matriculas,
                           selected_disciplina_id=disciplina_id,
                           selected_turma_id=turma_id,
                           selected_aluno_id=aluno_id,
                           selected_status=status)


@app.route("/relatorios/download/<format>", methods=["GET"])
@login_required
def download_relatorio(format):
    disciplina_id = request.args.get("disciplina_id", type=int) or None
    turma_id = request.args.get("turma_id", type=int) or None
    aluno_id = request.args.get("aluno_id", type=int) or None
    status = request.args.get("status", "").strip() or None

    matriculas = get_relatorio_data(disciplina_id, turma_id, aluno_id, status)

    if not matriculas:
        flash("Nenhum dado para gerar o relatório.", "aviso")
        return redirect(url_for("relatorios"))

    if format == "pdf":
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4))
        styles = getSampleStyleSheet()
        elements = []

        elements.append(Paragraph("Relatório de Matrículas", styles['h1']))
        elements.append(Spacer(1, 0.2 * inch))

        data = [['Aluno', 'Disciplina', 'Turma', 'Início', 'Conclusão', 'Média Final', 'Frequência (%)', 'Status']]
        for m in matriculas:
            data.append([
                m['aluno_nome'],
                m['disciplina_nome'],
                m['turma_nome'] or 'N/A',
                m['data_inicio'],
                m['data_conclusao'] or 'Cursando',
                m['media_display'],
                f"{m['frequencia_porcentagem']:.1f}%",
                m['status_display']
            ])

        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#007bff')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 2),
            ('RIGHTPADDING', (0, 0), (-1, -1), 2),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ])

        col_widths = [1.8*inch, 1.8*inch, 1.5*inch, 1*inch, 1*inch, 1*inch, 1.2*inch, 1.5*inch]
        table = Table(data, colWidths=col_widths)
        table.setStyle(table_style)
        elements.append(table)

        doc.build(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="relatorio_matriculas.pdf", mimetype="application/pdf")

    elif format == "docx":
        document = Document()
        document.add_heading('Relatório de Matrículas', level=1)

        table = document.add_table(rows=1, cols=8)
        table.autofit = True
        table.allow_autofit = True
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(1.5)
        table.columns[2].width = Inches(1.2)
        table.columns[3].width = Inches(0.8)
        table.columns[4].width = Inches(0.8)
        table.columns[5].width = Inches(0.8)
        table.columns[6].width = Inches(1.0)
        table.columns[7].width = Inches(1.2)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Aluno'
        hdr_cells[1].text = 'Disciplina'
        hdr_cells[2].text = 'Turma'
        hdr_cells[3].text = 'Início'
        hdr_cells[4].text = 'Conclusão'
        hdr_cells[5].text = 'Média Final'
        hdr_cells[6].text = 'Frequência (%)'
        hdr_cells[7].text = 'Status'

        # Estilo para o cabeçalho da tabela
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(0xFF, 0xFF, 0xFF) # Branco
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            # Adicionar sombreamento azul (não é direto, precisa de XML)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '007bff') # Cor azul
            cell._tc.get_or_add_tcPr().append(shading_elm)


        for m in matriculas:
            row_cells = table.add_row().cells
            row_cells[0].text = m['aluno_nome']
            row_cells[1].text = m['disciplina_nome']
            row_cells[2].text = m['turma_nome'] or 'N/A'
            row_cells[3].text = m['data_inicio']
            row_cells[4].text = m['data_conclusao'] or 'Cursando'
            row_cells[5].text = m['media_display']
            row_cells[6].text = f"{m['frequencia_porcentagem']:.1f}%"
            row_cells[7].text = m['status_display']
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(8)


        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="relatorio_matriculas.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    flash("Formato de relatório inválido.", "erro")
    return redirect(url_for("relatorios"))


@app.route("/relatorios/frequencia", methods=["GET", "POST"])
@login_required
def relatorios_frequencia():
    conn = conectar()
    cursor = conn.cursor()
    disciplinas = []
    turmas = []
    alunos = []
    frequencia_data = []
    selected_disciplina_id = None
    selected_turma_id = None
    selected_aluno_id = None
    selected_data_inicio = None
    selected_data_fim = None

    try:
        cursor.execute("SELECT id, nome FROM disciplinas WHERE ativa=1 ORDER BY nome")
        disciplinas = cursor.fetchall()
        cursor.execute("SELECT id, nome FROM turmas WHERE ativa=1 ORDER BY nome")
        turmas = cursor.fetchall()
        cursor.execute("SELECT id, nome FROM alunos ORDER BY nome")
        alunos = cursor.fetchall()

        if request.method == "POST":
            selected_disciplina_id = request.form.get("disciplina_id", type=int) or None
            selected_turma_id = request.form.get("turma_id", type=int) or None
            selected_aluno_id = request.form.get("aluno_id", type=int) or None
            selected_data_inicio = request.form.get("data_inicio", "").strip() or None
            selected_data_fim = request.form.get("data_fim", "").strip() or None

            query = """
                SELECT a.nome AS aluno_nome,
                       d.nome AS disciplina_nome,
                       t.nome AS turma_nome,
                       p.data_aula,
                       p.presente,
                       p.fez_atividade,
                       d.tem_atividades
                FROM presencas p
                JOIN matriculas m ON p.matricula_id = m.id
                JOIN alunos a ON m.aluno_id = a.id
                JOIN disciplinas d ON m.disciplina_id = d.id
                LEFT JOIN turmas t ON a.turma_id = t.id
                WHERE 1=1
            """
            params = []

            if selected_disciplina_id:
                query += " AND d.id = ?"
                params.append(selected_disciplina_id)
            if selected_turma_id:
                query += " AND t.id = ?"
                params.append(selected_turma_id)
            if selected_aluno_id:
                query += " AND a.id = ?"
                params.append(selected_aluno_id)
            if selected_data_inicio:
                query += " AND p.data_aula >= ?"
                params.append(selected_data_inicio)
            if selected_data_fim:
                query += " AND p.data_aula <= ?"
                params.append(selected_data_fim)

            query += " ORDER BY a.nome, d.nome, p.data_aula DESC"
            cursor.execute(query, params)
            frequencia_data = [dict(row) for row in cursor.fetchall()]

    except Exception as e:
        flash(f"Erro ao gerar relatório de frequência: {e}", "erro")
        print(f"ERRO NO RELATÓRIO DE FREQUÊNCIA: {e}") # Logar o erro no console
    finally:
        conn.close()

    return render_template("relatorio_frequencia.html",
                           disciplinas=disciplinas,
                           turmas=turmas,
                           alunos=alunos,
                           frequencia_data=frequencia_data,
                           selected_disciplina_id=selected_disciplina_id,
                           selected_turma_id=selected_turma_id,
                           selected_aluno_id=selected_aluno_id,
                           selected_data_inicio=selected_data_inicio,
                           selected_data_fim=selected_data_fim)


@app.route("/relatorios/frequencia/download/<format>", methods=["GET"])
@login_required
def download_relatorio_frequencia(format):
    selected_disciplina_id = request.args.get("disciplina_id", type=int) or None
    selected_turma_id = request.args.get("turma_id", type=int) or None
    selected_aluno_id = request.args.get("aluno_id", type=int) or None
    selected_data_inicio = request.args.get("data_inicio", "").strip() or None
    selected_data_fim = request.args.get("data_fim", "").strip() or None

    conn = conectar()
    cursor = conn.cursor()
    frequencia_data = []
    try:
        query = """
            SELECT a.nome AS aluno_nome,
                   d.nome AS disciplina_nome,
                   t.nome AS turma_nome,
                   p.data_aula,
                   p.presente,
                   p.fez_atividade,
                   d.tem_atividades
            FROM presencas p
            JOIN matriculas m ON p.matricula_id = m.id
            JOIN alunos a ON m.aluno_id = a.id
            JOIN disciplinas d ON m.disciplina_id = d.id
            LEFT JOIN turmas t ON a.turma_id = t.id
            WHERE 1=1
        """
        params = []

        if selected_disciplina_id:
            query += " AND d.id = ?"
            params.append(selected_disciplina_id)
        if selected_turma_id:
            query += " AND t.id = ?"
            params.append(selected_turma_id)
        if selected_aluno_id:
            query += " AND a.id = ?"
            params.append(selected_aluno_id)
        if selected_data_inicio:
            query += " AND p.data_aula >= ?"
            params.append(selected_data_inicio)
        if selected_data_fim:
            query += " AND p.data_aula <= ?"
            params.append(selected_data_fim)

        query += " ORDER BY a.nome, d.nome, p.data_aula DESC"
        cursor.execute(query, params)
        frequencia_data = [dict(row) for row in cursor.fetchall()]
    except Exception as e:
        flash(f"Erro ao buscar dados para download do relatório de frequência: {e}", "erro")
        print(f"ERRO NO DOWNLOAD DO RELATÓRIO DE FREQUÊNCIA: {e}")
        return redirect(url_for("relatorios_frequencia"))
    finally:
        conn.close()

    if not frequencia_data:
        flash("Nenhum dado para gerar o relatório de frequência.", "aviso")
        return redirect(url_for("relatorios_frequencia"))

    if format == "pdf":
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4))
        styles = getSampleStyleSheet()
        elements = []

        elements.append(Paragraph("Relatório de Frequência", styles['h1']))
        elements.append(Spacer(1, 0.2 * inch))

        data = [['Aluno', 'Disciplina', 'Turma', 'Data da Aula', 'Presente', 'Atividade Feita']]
        for f in frequencia_data:
            presente_str = "Sim" if f['presente'] else "Não"
            atividade_str = "Sim" if f['tem_atividades'] and f['fez_atividade'] else ("Não" if f['tem_atividades'] else "N/A")
            data.append([
                f['aluno_nome'],
                f['disciplina_nome'],
                f['turma_nome'] or 'N/A',
                f['data_aula'],
                presente_str,
                atividade_str
            ])

        table_style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#007bff')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('LEFTPADDING', (0, 0), (-1, -1), 2),
            ('RIGHTPADDING', (0, 0), (-1, -1), 2),
            ('TOPPADDING', (0, 0), (-1, -1), 2),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ])

        col_widths = [1.8*inch, 1.8*inch, 1.5*inch, 1.2*inch, 1*inch, 1.5*inch]
        table = Table(data, colWidths=col_widths)
        table.setStyle(table_style)
        elements.append(table)

        doc.build(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="relatorio_frequencia.pdf", mimetype="application/pdf")

    elif format == "docx":
        document = Document()
        document.add_heading('Relatório de Frequência', level=1)

        table = document.add_table(rows=1, cols=6)
        table.autofit = True
        table.allow_autofit = True
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(1.5)
        table.columns[2].width = Inches(1.2)
        table.columns[3].width = Inches(1.0)
        table.columns[4].width = Inches(0.8)
        table.columns[5].width = Inches(1.2)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Aluno'
        hdr_cells[1].text = 'Disciplina'
        hdr_cells[2].text = 'Turma'
        hdr_cells[3].text = 'Data da Aula'
        hdr_cells[4].text = 'Presente'
        hdr_cells[5].text = 'Atividade Feita'

        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.color.rgb = docx.shared.RGBColor(0xFF, 0xFF, 0xFF)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), '007bff')
            cell._tc.get_or_add_tcPr().append(shading_elm)

        for f in frequencia_data:
            row_cells = table.add_row().cells
            presente_str = "Sim" if f['presente'] else "Não"
            atividade_str = "Sim" if f['tem_atividades'] and f['fez_atividade'] else ("Não" if f['tem_atividades'] else "N/A")
            row_cells[0].text = f['aluno_nome']
            row_cells[1].text = f['disciplina_nome']
            row_cells[2].text = f['turma_nome'] or 'N/A'
            row_cells[3].text = f['data_aula']
            row_cells[4].text = presente_str
            row_cells[5].text = atividade_str
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].runs[0].font.size = Pt(8)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="relatorio_frequencia.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    flash("Formato de relatório inválido.", "erro")
    return redirect(url_for("relatorios_frequencia"))


# ══════════════════════════════════════
# USUÁRIOS
# ══════════════════════════════════════
@app.route("/usuarios")
@login_required
@admin_required
def usuarios():
    conn   = conectar()
    cursor = conn.cursor()
    cursor.execute("SELECT id, nome, email, perfil FROM usuarios ORDER BY nome")
    lista = cursor.fetchall()
    conn.close()
    return render_template("usuarios.html", usuarios=lista)


@app.route("/usuarios/novo", methods=["GET", "POST"])
@login_required
@admin_required
def novo_usuario():
    if request.method == "POST":
        nome    = request.form.get("nome", "").strip()
        email   = request.form.get("email", "").strip()
        senha   = request.form.get("senha", "")
        perfil  = request.form.get("perfil", "").strip()

        if not nome or not email or not senha or not perfil:
            flash("Todos os campos são obrigatórios!", "erro")
            return render_template("novo_usuario.html")
        if len(senha) < 6:
            flash("A senha deve ter no mínimo 6 caracteres!", "erro")
            return render_template("novo_usuario.html")

        conn   = conectar()
        cursor = conn.cursor()
        try:
            senha_hash = generate_password_hash(senha)
            cursor.execute(
                "INSERT INTO usuarios (nome, email, senha_hash, perfil) VALUES (?, ?, ?, ?)",
                (nome, email, senha_hash, perfil))
            conn.commit()
            flash(f"Usuário '{nome}' criado!", "sucesso")
        except sqlite3.IntegrityError:
            flash("Já existe um usuário com este e-mail!", "erro")
        except Exception as e:
            flash(f"Erro ao criar usuário: {e}", "erro")
        finally:
            conn.close()
        return redirect(url_for("usuarios"))
    return render_template("novo_usuario.html")


@app.route("/usuarios/<int:id>/editar", methods=["GET", "POST"])
@login_required
@admin_required
def editar_usuario(id):
    conn   = conectar()
    cursor = conn.cursor()
    usuario = None
    try:
        cursor.execute("SELECT id, nome, email, perfil FROM usuarios WHERE id=?", (id,))
        usuario = cursor.fetchone()
    except Exception as e:
        flash(f"Erro ao carregar usuário: {e}", "erro")
    finally:
        conn.close()

    if not usuario:
        flash("Usuário não encontrado!", "erro")
        return redirect(url_for("usuarios"))

    if request.method == "POST":
        nome    = request.form.get("nome", "").strip()
        email   = request.form.get("email", "").strip()
        perfil  = request.form.get("perfil", "").strip()
        senha   = request.form.get("senha", "").strip() # Senha opcional para alteração

        if not nome or not email or not perfil:
            flash("Nome, e-mail e perfil são obrigatórios!", "erro")
            return render_template("editar_usuario.html", usuario=usuario)

        conn = conectar()
        cursor = conn.cursor()
        try:
            if senha: # Se uma nova senha foi fornecida
                if len(senha) < 6:
                    flash("A nova senha deve ter no mínimo 6 caracteres!", "erro")
                    return render_template("editar_usuario.html", usuario=usuario)
                senha_hash = generate_password_hash(senha)
                cursor.execute(
                    "UPDATE usuarios SET nome=?, email=?, perfil=?, senha_hash=? WHERE id=?",
                    (nome, email, perfil, senha_hash, id))
            else:
                cursor.execute(
                    "UPDATE usuarios SET nome=?, email=?, perfil=? WHERE id=?",
                    (nome, email, perfil, id))
            conn.commit()
            flash(f"Usuário '{nome}' atualizado!", "sucesso")
        except sqlite3.IntegrityError:
            flash("Já existe um usuário com este e-mail!", "erro")
        except Exception as e:
            flash(f"Erro ao atualizar usuário: {e}", "erro")
        finally:
            conn.close()
        return redirect(url_for("usuarios"))
    return render_template("editar_usuario.html", usuario=usuario)


@app.route("/usuarios/<int:id>/excluir", methods=["POST"])
@login_required
@admin_required
def excluir_usuario(id):
    if id == current_user.id:
        flash("Você não pode excluir sua própria conta!", "erro")
        return redirect(url_for("usuarios"))

    conn   = conectar()
    cursor = conn.cursor()
    try:
        cursor.execute("DELETE FROM usuarios WHERE id=?", (id,))
        conn.commit()
        flash("Usuário excluído!", "sucesso")
    except Exception as e:
        flash(f"Erro ao excluir usuário: {e}", "erro")
    finally:
        conn.close()
    return redirect(url_for("usuarios"))


@app.route("/minha-conta", methods=["GET", "POST"])
@login_required
def minha_conta():
    if request.method == "POST":
        senha_atual = request.form.get("senha_atual", "")
        nova_senha  = request.form.get("nova_senha", "")
        confirmar   = request.form.get("confirmar", "")
        u = verificar_login(current_user.email, senha_atual)
        if not u:
            flash("Senha atual incorreta!", "erro")
            return redirect(url_for("minha_conta"))
        if nova_senha != confirmar:
            flash("As senhas não coincidem!", "erro")
            return redirect(url_for("minha_conta"))
        if len(nova_senha) < 6:
            flash("Mínimo 6 caracteres!", "erro")
            return redirect(url_for("minha_conta"))
        conn   = conectar()
        cursor = conn.cursor()
        try:
            cursor.execute(
                "UPDATE usuarios SET senha_hash=? WHERE id=?",
                (generate_password_hash(nova_senha), current_user.id))
            conn.commit()
            flash("Senha alterada com sucesso!", "sucesso")
        except Exception as e:
            flash(f"Erro ao alterar senha: {e}", "erro")
        finally:
            conn.close()
        return redirect(url_for("index"))
    return render_template("minha_conta.html")


# ══════════════════════════════════════
# BACKUP E RESTAURAÇÃO
# ══════════════════════════════════════
@app.route("/admin/backup", methods=["GET", "POST"])
@login_required
@admin_required
def backup_restauracao():
    if request.method == "POST":
        if 'backup_action' in request.form: # Requisição para baixar backup
            try:
                backup_filename = f"escola_backup_{date.today().strftime('%Y-%m-%d')}.db"
                return send_file(DATABASE, as_attachment=True, download_name=backup_filename, mimetype="application/x-sqlite3")
            except Exception as e:
                flash(f"Erro ao gerar backup: {e}", "erro")
                print(f"ERRO NO BACKUP: {e}")
        elif 'restore_file' in request.files: # Requisição para restaurar
            file = request.files['restore_file']
            if file and file.filename.endswith('.db'):
                try:
                    # Fechar a conexão atual com o banco de dados antes de substituir o arquivo
                    # Isso é crucial para evitar erros de arquivo em uso
                    # (A função conectar() cria uma nova conexão, então não precisamos fechar explicitamente aqui)

                    # Criar um backup temporário do banco de dados atual antes de sobrescrever
                    temp_backup_path = DATABASE + ".temp_bak"
                    shutil.copy(DATABASE, temp_backup_path)

                    file.save(DATABASE)
                    flash("Banco de dados restaurado com sucesso! Pode ser necessário reiniciar o servidor para que as mudanças sejam totalmente aplicadas.", "sucesso")
                    # Remover o backup temporário após sucesso
                    if os.path.exists(temp_backup_path):
                        os.remove(temp_backup_path)
                except Exception as e:
                    # Se houver erro, tentar restaurar do backup temporário
                    if os.path.exists(temp_backup_path):
                        shutil.copy(temp_backup_path, DATABASE)
                        os.remove(temp_backup_path) # Remover o temp_bak
                    flash(f"Erro ao restaurar banco de dados: {e}", "erro")
                    print(f"ERRO NA RESTAURAÇÃO: {e}")
            else:
                flash("Por favor, selecione um arquivo de backup válido (.db).", "erro")
        else:
            flash("Ação inválida.", "erro")
        return redirect(url_for("backup_restauracao"))
    return render_template("backup_restauracao.html")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)